import os
from flask import Flask, send_file, request
import pandas as pd
from datetime import datetime
from io import BytesIO
from docx import Document

app = Flask(__name__)

# Путь к файлу Excel
EXCEL_FILE = 'phishing_data.xlsx'

# Удаляем файл, если он поврежден
def reset_excel_file():
    if os.path.exists(EXCEL_FILE):
        try:
            # Попытка открыть файл как Excel
            df = pd.read_excel(EXCEL_FILE, engine='openpyxl')
        except Exception as e:
            print(f"Ошибка при открытии файла: {e}. Удаляем файл и создаем новый.")
            os.remove(EXCEL_FILE)

# Инициализация файла Excel, если он не существует
def initialize_excel():
    reset_excel_file()  # Удаляем поврежденный файл, если он существует
    if not os.path.exists(EXCEL_FILE):
        try:
            # Создаем новый DataFrame с нужными столбцами
            df = pd.DataFrame(columns=["Timestamp", "IP Address"])
            df.to_excel(EXCEL_FILE, index=False, engine='openpyxl')
            print("Новый Excel файл создан.")
        except Exception as e:
            print(f"Ошибка при создании Excel файла: {e}")

# Функция для записи данных о переходе
def log_redirect(ip_address):
    try:
        timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
        
        # Чтение и добавление данных в Excel
        df = pd.read_excel(EXCEL_FILE, engine='openpyxl')
        
        # Создаем новый DataFrame с одной строкой данных
        new_data = pd.DataFrame({"Timestamp": [timestamp], "IP Address": [ip_address]})
        
        # Используем pd.concat для добавления новых данных в DataFrame
        df = pd.concat([df, new_data], ignore_index=True)
        
        # Сохраняем измененный DataFrame обратно в Excel
        df.to_excel(EXCEL_FILE, index=False, engine='openpyxl')
        print(f"Логирование IP: {ip_address} успешно!")
    except Exception as e:
        print(f"Ошибка при записи в Excel: {e}")

# Функция для создания Word-документа
def create_word_file():
    doc = Document()
    doc.add_heading('Проверка безопасности', 0)

    # Добавление текста в документ
    doc.add_paragraph("Проверка не пройдена")

    # Сохраняем документ в памяти (не на диск)
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)  # Перемещаем курсор в начало файла для отправки
    return byte_io

# Редирект на реальный сайт и отправка файла
@app.route('/')
def phishing_redirect():
    try:
        ip_address = request.remote_addr  # Получаем IP-адрес
        log_redirect(ip_address)  # Логируем информацию

        # Создаем Word-документ
        word_file = create_word_file()

        # Отправляем файл для скачивания
        return send_file(word_file, as_attachment=True, download_name="ваши_файлы_будут_зашифрованы_через_минуту.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    except Exception as e:
        print(f"Ошибка при обработке запроса: {e}")
        return "Произошла ошибка на сервере. Пожалуйста, попробуйте позже.", 500

if __name__ == '__main__':
    try:
        initialize_excel()  # Инициализация файла
        app.run(host='0.0.0.0', port=8080)
    except Exception as e:
        print(f"Ошибка при запуске сервера: {e}")
